#!/usr/bin/env python3
"""Export a clean reading DOCX from the annotated Markdown manuscript."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError as exc:  # pragma: no cover - environment guidance
    raise SystemExit("缺少 python-docx；请使用 Codex 工作区 Python 运行此脚本。") from exc


EPISODE = re.compile(r"^##\s+(第\s*[一二三四五六七八九十百千万0-9]+\s*集[：:].+?)\s*$", re.MULTILINE)
BODY_HEADING = "### 本集正文"
FORBIDDEN_WORKING_LABELS = ("本集创作标注", "本集正文", "开头钩子类型", "本集正文字符数")


def clean_inline_markdown(text: str) -> str:
    text = re.sub(r"!\[[^]]*]\([^)]+\)", "", text)
    text = re.sub(r"\[([^]]+)]\([^)]+\)", r"\1", text)
    text = re.sub(r"(?<!\\)(?:\*\*|__|\*|_)(.+?)(?:\*\*|__|\*|_)", r"\1", text)
    return text.replace("\\*", "*").strip()


def extract_episodes(markdown: str) -> list[tuple[str, list[str]]]:
    matches = list(EPISODE.finditer(markdown))
    episodes: list[tuple[str, list[str]]] = []
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(markdown)
        block = markdown[match.end():end]
        body_pos = block.find(BODY_HEADING)
        if body_pos < 0:
            raise ValueError(f"{match.group(1)} 缺少“{BODY_HEADING}”")
        body = block[body_pos + len(BODY_HEADING):]
        lines = [clean_inline_markdown(line) for line in body.splitlines()]
        while lines and not lines[0]:
            lines.pop(0)
        while lines and not lines[-1]:
            lines.pop()
        if not any(lines):
            raise ValueError(f"{match.group(1)} 的正文为空")
        episodes.append((clean_inline_markdown(match.group(1)), lines))
    if not episodes:
        raise ValueError("未找到“## 第 N 集：标题”格式的分集")
    return episodes


def set_east_asia_font(style, name: str, size: float) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def export_docx(manuscript: Path, output: Path, title: str) -> int:
    episodes = extract_episodes(manuscript.read_text(encoding="utf-8"))
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)

    set_east_asia_font(document.styles["Normal"], "宋体", 11)
    set_east_asia_font(document.styles["Title"], "黑体", 22)
    set_east_asia_font(document.styles["Heading 1"], "黑体", 16)
    document.styles["Normal"].paragraph_format.line_spacing = 1.5
    document.styles["Normal"].paragraph_format.space_after = Pt(6)

    title_paragraph = document.add_paragraph(style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.add_run(title)

    for index, (episode_title, lines) in enumerate(episodes):
        if index:
            document.add_page_break()
        heading = document.add_paragraph(episode_title, style="Heading 1")
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for line in lines:
            if not line:
                continue
            paragraph = document.add_paragraph(line)
            paragraph.paragraph_format.first_line_indent = Cm(0.74)

    document.core_properties.title = title
    document.core_properties.subject = "AI 漫剧小说正式阅读稿"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return len(episodes)


def main() -> None:
    parser = argparse.ArgumentParser(description="把带创作标注的 Markdown 导出为干净的 Word 阅读稿。")
    parser.add_argument("manuscript", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--title", default="漫剧小说正式稿")
    args = parser.parse_args()
    try:
        count = export_docx(args.manuscript, args.output, args.title)
    except ValueError as exc:
        raise SystemExit(f"导出失败：{exc}") from exc
    print(f"导出完成：{args.output}；共 {count} 集。")


if __name__ == "__main__":
    main()
